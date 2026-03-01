"""
Test if python-docx loses shapes just by opening and saving a document
without any modifications.
"""

import os
from docx import Document


def count_shapes(doc):
    """Count all shape/drawing elements in a document using comprehensive detection."""
    count = 0
    shape_tags = set()
    
    for para in doc.paragraphs:
        for elem in para._element.iter():
            tag = elem.tag.lower()
            if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                count += 1
                shape_tags.add(elem.tag)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for elem in para._element.iter():
                        tag = elem.tag.lower()
                        if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                            count += 1
                            shape_tags.add(elem.tag)
    
    for section in doc.sections:
        for para in section.header.paragraphs:
            for elem in para._element.iter():
                tag = elem.tag.lower()
                if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                    count += 1
                    shape_tags.add(elem.tag)
        
        for para in section.footer.paragraphs:
            for elem in para._element.iter():
                tag = elem.tag.lower()
                if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                    count += 1
                    shape_tags.add(elem.tag)
    
    return count, list(shape_tags)


def main():
    source_file = r"C:\Users\theam\Downloads\medical-certificate-template-09\medical certificate template_es.docx"
    test_output = r"d:\PROJECTS\AI TRANSLATOR\test_output.docx"
    
    print("="*70)
    print("PYTHON-DOCX SHAPE PRESERVATION TEST")
    print("="*70)
    print("\nTest: Open document, change nothing, save it\n")
    
    # Load original
    print(f"Loading original: {os.path.basename(source_file)}")
    original_doc = Document(source_file)
    original_shapes, shape_tags = count_shapes(original_doc)
    print(f"Original shapes: {original_shapes}")
    if shape_tags:
        print(f"Shape tags found: {shape_tags[:5]}...")  # Show first 5
    
    # Save without modification
    print(f"\nSaving without any modifications to: {os.path.basename(test_output)}")
    original_doc.save(test_output)
    
    # Load the saved document
    print("Reloading the saved document...")
    saved_doc = Document(test_output)
    saved_shapes, _ = count_shapes(saved_doc)
    print(f"Saved document shapes: {saved_shapes}")
    
    # Compare
    print("\n" + "="*70)
    if original_shapes == saved_shapes:
        print(f"✓ RESULT: Shapes preserved correctly ({saved_shapes} shapes)")
    else:
        print(f"✗ RESULT: Shapes lost during open/save!")
        print(f"  Original: {original_shapes}, Saved: {saved_shapes}")
        print(f"  Lost: {original_shapes - saved_shapes} shapes")
    print("="*70)


if __name__ == "__main__":
    main()
