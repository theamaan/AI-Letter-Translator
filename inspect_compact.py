"""Compact inspection of the example .docx file."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
 
path = r'C:\Users\theam\Downloads\medical-certificate-template-09'
filename = '453-medical-professional-blue-portrait_es.docx'
doc = Document(os.path.join(path, filename))
 
print(f'=== DOCUMENT: {filename} ===')
print(f'Sections: {len(doc.sections)}, Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
print()
 
# Just show paragraph text with key formatting info (compact)
print('=== ALL PARAGRAPHS (text only, with style/alignment) ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text if para.text else '[EMPTY]'
    style = para.style.name if para.style else '-'
   
    # Check if runs have mixed formatting
    formats = set()
    for run in para.runs:
        f = []
        if run.font.bold: f.append('B')
        if run.font.italic: f.append('I')
        if run.font.underline: f.append('U')
        formats.add(''.join(f) if f else 'normal')
    fmt_str = '+'.join(sorted(formats)) if formats else '-'
   
    num_runs = len(para.runs)
    print(f'P{i:03d} [{style:15s}] runs={num_runs:3d} fmt={fmt_str:10s} | {text[:200]}')
 
# Tables
print('\n=== TABLES ===')
for t_idx, table in enumerate(doc.tables):
    print(f'\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols')
    for r_idx, row in enumerate(table.rows):
        row_text = ' | '.join(cell.text[:60] for cell in row.cells)
        print(f'  Row{r_idx}: {row_text}')
 
# Headers/Footers
print('\n=== HEADERS/FOOTERS ===')
for s_idx, section in enumerate(doc.sections):
    for p in section.header.paragraphs:
        if p.text.strip():
            print(f'  Header[{s_idx}]: {p.text[:150]}')
    for p in section.footer.paragraphs:
        if p.text.strip():
            print(f'  Footer[{s_idx}]: {p.text[:150]}')
 
# Images / inline shapes
print(f'\n=== INLINE SHAPES: {len(doc.inline_shapes)} ===')
