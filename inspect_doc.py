"""Inspect a .docx file to understand its structure and formatting."""
from docx import Document
import os
 
path = r'D:\Python Project\Input Letters'
filename = 'medical certificate template_es.docx'
doc = Document(os.path.join(path, filename))
 
print('=== DOCUMENT STRUCTURE ===')
print(f'Sections: {len(doc.sections)}')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
print()
 
# Show paragraph details with formatting
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else 'None'
    align = str(para.alignment) if para.alignment else 'None'
    text = para.text[:150] if para.text else '[EMPTY]'
    print(f'P{i:03d} | Style={style:20s} | Align={align:30s} | Text: {text}')
    # Show runs for non-empty paragraphs
    if para.text.strip():
        for j, run in enumerate(para.runs):
            font = run.font
            bold = font.bold
            italic = font.italic
            size = font.size
            name = font.name
            color = font.color.rgb if font.color and font.color.rgb else None
            print(f'  Run{j}: bold={bold}, italic={italic}, size={size}, font={name}, color={color}')
            print(f'    text="{run.text[:100]}"')
 
# Show tables
print('\n=== TABLES ===')
for t_idx, table in enumerate(doc.tables):
    print(f'\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols')
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f'  [{r_idx},{c_idx}]: {cell.text[:100]}')
 
# Check headers/footers
print('\n=== HEADERS/FOOTERS ===')
for s_idx, section in enumerate(doc.sections):
    header = section.header
    footer = section.footer
    if header and not header.is_linked_to_previous:
        for p in header.paragraphs:
            if p.text.strip():
                print(f'  Header S{s_idx}: {p.text[:100]}')
    if footer and not footer.is_linked_to_previous:
        for p in footer.paragraphs:
            if p.text.strip():
                print(f'  Footer S{s_idx}: {p.text[:100]}')