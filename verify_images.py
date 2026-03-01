"""
Enhanced verification script to check all types of images in Word documents.
"""

import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def count_all_image_types(doc, doc_path):
    """Count all image/shape elements in a Word document."""
    shape_count = 0
    image_count = 0
    
    # Check for images in document.xml relationships
    try:
        doc_part = doc.part
        if hasattr(doc_part, 'related_parts'):
            for rel_id, rel in doc_part.related_parts.items():
                if 'image' in rel.reltype.lower():
                    image_count += 1
    except:
        pass
    
    # Check all shape-related tags using comprehensive detection
    for para in doc.paragraphs:
        for elem in para._element.iter():
            tag = elem.tag.lower()
            if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                shape_count += 1
    
    # Check in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for elem in para._element.iter():
                        tag = elem.tag.lower()
                        if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                            shape_count += 1
    
    # Check in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for elem in para._element.iter():
                tag = elem.tag.lower()
                if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                    shape_count += 1
        
        for para in section.footer.paragraphs:
            for elem in para._element.iter():
                tag = elem.tag.lower()
                if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
                    shape_count += 1
    
    return image_count, shape_count


def main():
    source_file = r"C:\Users\theam\Downloads\medical-certificate-template-09\medical certificate template_es.docx"
    translated_file = r"d:\PROJECTS\AI TRANSLATOR\Translated Files\medical certificate template_es.docx"
    
    print("="*70)
    print("COMPREHENSIVE IMAGE/LOGO PRESERVATION VERIFICATION")
    print("="*70)
    
    # Check if files exist
    if not os.path.exists(source_file):
        print(f"\nERROR: Source file not found")
        print(f"  Path: {source_file}")
        return
    
    if not os.path.exists(translated_file):
        print(f"\nERROR: Translated file not found")
        print(f"  Path: {translated_file}")
        return
    
    # Load documents
    print("\nLoading documents...")
    source_doc = Document(source_file)
    translated_doc = Document(translated_file)
    
    # Count relationships
    source_img_rels = count_all_image_types(source_doc, source_file)[0]
    source_shapes = count_all_image_types(source_doc, source_file)[1]
    
    translated_img_rels = count_all_image_types(translated_doc, translated_file)[0]
    translated_shapes = count_all_image_types(translated_doc, translated_file)[1]
    
    print("\n" + "-"*70)
    print("RESULTS:")
    print("-"*70)
    print(f"\nSource document:")
    print(f"  Image resources (via relationships): {source_img_rels}")
    print(f"  Shape/drawing elements: {source_shapes}")
    print(f"  Total: {source_img_rels + source_shapes}")
    
    print(f"\nTranslated document:")
    print(f"  Image resources (via relationships): {translated_img_rels}")
    print(f"  Shape/drawing elements: {translated_shapes}")
    print(f"  Total: {translated_img_rels + translated_shapes}")
    
    # Verify preservation
    print("\n" + "-"*70)
    source_total = source_img_rels + source_shapes
    translated_total = translated_img_rels + translated_shapes
    
    if source_total == 0:
        print("ℹ No images/logos detected in source document")
    elif source_total == translated_total:
        print(f"✓ SUCCESS: All {source_total} images/logos were preserved!")
    else:
        print(f"✗ ISSUE: Image count mismatch!")
        print(f"  Expected: {source_total}, Found: {translated_total}")
        missing = source_total - translated_total
        if missing > 0:
            print(f"  Missing: {missing} images/logos")
    print("-"*70)


if __name__ == "__main__":
    main()
