"""
AI Letter Translator — Agentic Document Translation Tool
=========================================================
Reads .docx healthcare letters, translates them using a hybrid LLM strategy:

  - Ollama (local open-source model) for document analysis and simple translations
  - Cerebras API (premium model) for complex medical/legal content and fallback

Optimisations over the single-model version:
  - Prompt compression   : ~60 % fewer input tokens per translation call
  - Paragraph batching   : up to 5 short paragraphs per API call
  - Per-document cache   : repeated phrases cost zero tokens
  - Smart DNT inclusion  : do-not-translate list omitted when a paragraph
                           contains no protected items

Usage:
    python translator.py                               # Translate first file found
    python translator.py --file "filename.docx"        # Translate a specific file
    python translator.py --all                         # Translate all files
    python translator.py --provider hybrid             # Auto-routing (default)
    python translator.py --provider ollama-only        # Force local model only
    python translator.py --provider cerebras-only      # Force Cerebras only
    python translator.py --all --skip-existing         # Skip already translated

Author: AI Translator Automation
"""

import argparse
import copy
import json
import os
import re
import sys
import time
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Emu, RGBColor

import config
from routing import TranslationRouter
from spanish_readability import (
    SpanishReadabilityTracker,
    SpanishTargetBand,
    adapt_spanish_text_to_grade,
    normalize_target_grade,
)


# ---------------------------------------------------------------------------
# Translation Stats
# ---------------------------------------------------------------------------

@dataclass
class _TranslationStats:
    """Tracks per-document API usage so cost-saving metrics can be reported."""
    cerebras_calls: int = 0
    ollama_calls: int   = 0
    fallback_calls: int = 0
    cache_hits: int     = 0
    batched_calls: int  = 0

    @property
    def total_api_calls(self) -> int:
        return self.cerebras_calls + self.ollama_calls

    def record(self, label: str) -> None:
        """Increment the appropriate counter from a provider label string."""
        low = label.lower()
        if "fallback" in low:
            self.fallback_calls += 1
            self.cerebras_calls += 1
        elif "cerebras" in low:
            self.cerebras_calls += 1
        else:
            self.ollama_calls += 1


# ---------------------------------------------------------------------------
# Backward-Compatible Shims
# ---------------------------------------------------------------------------
# These allow existing test scripts (test_api.py, test_ssl.py) to continue
# working without any modifications.

def get_client():
    """Return a CerebrasProvider instance (backward-compat shim)."""
    from llm_providers import CerebrasProvider
    return CerebrasProvider()


def call_llm(client, system_prompt: str, user_prompt: str) -> str:
    """
    Call an LLM provider and return the response text.
    Accepts a LLMProvider instance (new path) or a raw Cerebras SDK client (legacy).
    """
    from llm_providers import LLMProvider
    if isinstance(client, LLMProvider):
        return client.call(system_prompt, user_prompt)
    # Legacy raw Cerebras SDK client
    for attempt in range(1, config.MAX_RETRIES + 1):
        try:
            stream = client.chat.completions.create(
                model=config.CEREBRAS_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                stream=True,
                max_completion_tokens=20_000,
                temperature=0.1,
                top_p=1,
            )
            result = ""
            for chunk in stream:
                result += chunk.choices[0].delta.content or ""
            return result.strip()
        except Exception as exc:
            wait = config.RETRY_DELAY ** attempt
            print(f"  [Retry {attempt}/{config.MAX_RETRIES}] API error: {exc}. Waiting {wait}s...")
            time.sleep(wait)
    print("  ERROR: All API retries exhausted.")
    return ""


# ---------------------------------------------------------------------------
# Filename Parsing
# ---------------------------------------------------------------------------

def extract_language_code(filename: str) -> Optional[str]:
    """
    Extract the language code from the filename.
    E.g., 'MHWI Deny Your Request - OON_Member_..._es.docx' -> 'es'
    """
    name_without_ext = os.path.splitext(filename)[0]
    parts = name_without_ext.split("_")
    if parts:
        candidate = parts[-1].lower()
        if candidate in config.LANGUAGE_MAP:
            return candidate
    return None


def get_language_name(code: str) -> str:
    """Convert a language code to its full display name."""
    return config.LANGUAGE_MAP.get(code, code)


# ---------------------------------------------------------------------------
# Document Structure Helpers
# ---------------------------------------------------------------------------

def consolidate_runs(paragraph):
    """
    Consolidate adjacent runs with identical formatting into groups.
    Returns a list of dicts: [{text, run_indices, format_key}, ...]
    """
    if not paragraph.runs:
        return []

    groups = []
    current_group = {
        "text": paragraph.runs[0].text,
        "run_indices": [0],
        "format_key": _run_format_key(paragraph.runs[0]),
    }

    for i in range(1, len(paragraph.runs)):
        run = paragraph.runs[i]
        fmt_key = _run_format_key(run)
        if fmt_key == current_group["format_key"]:
            current_group["text"] += run.text
            current_group["run_indices"].append(i)
        else:
            groups.append(current_group)
            current_group = {
                "text": run.text,
                "run_indices": [i],
                "format_key": fmt_key,
            }
    groups.append(current_group)
    return groups


def _run_format_key(run) -> str:
    """Generate a hashable formatting key for a run."""
    font = run.font
    return "|".join([
        str(font.bold),
        str(font.italic),
        str(font.underline),
        str(font.size),
        str(font.name),
        str(font.color.rgb if font.color and font.color.rgb else None),
    ])


def extract_paragraph_text(paragraph) -> str:
    """Get the full text of a paragraph from all its runs."""
    return "".join(run.text for run in paragraph.runs)


def extract_paragraph_elements(paragraph):
    """
    Extract all elements (runs and inline shapes) from a paragraph in order.
    Returns a list of dicts: {type: 'run'|'shape', content: run|shape}
    """
    elements = []
    for elem in paragraph._element:
        if elem.tag.endswith('}r'):
            for run in paragraph.runs:
                if run._element == elem:
                    elements.append({'type': 'run', 'content': run})
                    break
        elif elem.tag.endswith('}pict') or elem.tag.endswith('}drawing'):
            elements.append({'type': 'shape', 'content': copy.deepcopy(elem)})
    return elements


def preserve_inline_shapes(paragraph):
    """
    Save a snapshot of all inline shape elements and their positions.
    Returns (shapes, positions) for later restoration.
    """
    shapes, positions = [], []
    for i, elem in enumerate(paragraph._element):
        if elem.tag.endswith('}pict') or elem.tag.endswith('}drawing'):
            shapes.append(copy.deepcopy(elem))
            positions.append(i)
    return shapes, positions


def restore_inline_shapes(paragraph, shapes, shape_positions):
    """Restore inline shapes to a paragraph at their original positions."""
    if not shapes:
        return

    current_elements = list(paragraph._element)
    for i in range(len(current_elements) - 1, -1, -1):
        elem = current_elements[i]
        if elem.tag.endswith('}pict') or elem.tag.endswith('}drawing'):
            paragraph._element.remove(elem)

    for shape, original_pos in zip(shapes, shape_positions):
        try:
            current_len = len(paragraph._element)
            if original_pos <= current_len:
                paragraph._element.insert(original_pos, shape)
            else:
                paragraph._element.append(shape)
        except Exception:
            paragraph._element.append(shape)


def _paragraph_has_shapes(paragraph) -> bool:
    """Return True if the paragraph contains any inline shapes or drawings."""
    for elem in paragraph._element.iter():
        tag = elem.tag.lower()
        if any(x in tag for x in ['shape', 'pict', 'drawing', 'blip', 'graphic', 'pic']):
            return True
    return False


# ---------------------------------------------------------------------------
# Optimised Prompts
# ---------------------------------------------------------------------------
# ~60% fewer tokens than the originals, reducing cost on every API call.

# Analysis prompt: ~300 chars vs. ~600 chars original.
_ANALYSIS_SYSTEM_PROMPT = (
    "Healthcare translation analyst. Identify ALL items that must NOT be translated: "
    "person names, street addresses, member/reference/case IDs, phone/fax numbers, "
    "emails, URLs, org names (e.g. \"Molina Healthcare\", \"BadgerCare Plus\"), "
    "medical codes (CPT/ICD/HCPCS), formatted dates, acronyms (TTY, TDD).\n\n"
    "Return ONLY this JSON, nothing else:\n"
    "{\"do_not_translate\": [\"item1\", ...], \"reasoning\": \"brief note\"}"
)

# Fixed tail appended to every translation prompt (short + constant = cache-friendly).
_TRANSLATION_FIXED_RULES = (
    "Also keep unchanged: person names, \"Molina Healthcare\", \"BadgerCare Plus\", "
    "phone/fax, emails, URLs, member IDs, medical codes, state abbrevs, ZIP codes.\n"
    "Return ONLY the translated text. No notes, no extra quotes."
)


def _contains_any_dnt(text: str, do_not_translate: list) -> bool:
    """Return True if 'text' contains at least one protected item."""
    text_lower = text.lower()
    return any(item.lower() in text_lower for item in do_not_translate)


def _build_translation_prompt(
    target_language: str,
    do_not_translate: list,
    text: str = "",
) -> str:
    """
    Build an optimised per-paragraph translation system prompt.
    DNT list is omitted when the paragraph contains no protected items,
    saving 50-100 tokens on the majority of short paragraphs.
    """
    lines = [
        f"Translate to {target_language}. Healthcare letter - exact meaning, formal tone.",
    ]
    if do_not_translate and (not text or _contains_any_dnt(text, do_not_translate)):
        dnt_compact = json.dumps(do_not_translate[:30], ensure_ascii=False)
        lines.append(f"KEEP UNCHANGED: {dnt_compact}")
    lines.append(_TRANSLATION_FIXED_RULES)
    return "\n".join(lines)


def _build_batch_system_prompt(target_language: str, do_not_translate: list) -> str:
    """System prompt for batch translation of multiple short paragraphs."""
    dnt_compact = (
        json.dumps(do_not_translate[:30], ensure_ascii=False)
        if do_not_translate else "[]"
    )
    return (
        f"Translate each numbered segment to {target_language}. "
        "Preserve the [N] markers exactly and output only the translated segments.\n"
        f"KEEP UNCHANGED: {dnt_compact}\n"
        + _TRANSLATION_FIXED_RULES
    )


def _parse_batch_response(response: str, count: int, originals: list) -> list:
    """Parse a [1]/[2]/[3]... batch response into individual translations."""
    results = []
    for idx in range(1, count + 1):
        pattern = rf"\[{idx}\]\s*(.*?)(?=\[{idx + 1}\]|\Z)"
        match = re.search(pattern, response, re.DOTALL)
        results.append(match.group(1).strip() if match else originals[idx - 1])
    return results


# ---------------------------------------------------------------------------
# Agentic Analysis
# ---------------------------------------------------------------------------

def analyze_document_content(router: TranslationRouter, full_text: str) -> tuple:
    """
    Agent Step 1: Identify elements that must NOT be translated.
    Routes to Ollama in hybrid mode; falls back to Cerebras automatically.

    Returns:
        (do_not_translate_list, provider_label)
    """
    print("  [Agent Step 1] Analysing document for non-translatable elements...")
    analysis_text = full_text[:6000]
    response, label = router.call_with_fallback(
        "analyze", _ANALYSIS_SYSTEM_PROMPT, analysis_text
    )
    try:
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
            cleaned = re.sub(r"\s*```$", "", cleaned)
        data = json.loads(cleaned)
        items = data.get("do_not_translate", [])
        reasoning = data.get("reasoning", "")
        print(f"  [Agent Step 1] Found {len(items)} non-translatable items via {label}.")
        if reasoning:
            print(f"  [Agent Step 1] {reasoning}")
        return items, label
    except (json.JSONDecodeError, KeyError) as exc:
        print(f"  [Agent Step 1] Warning: Could not parse analysis ({exc}). Using regex fallback.")
        return _fallback_extract_non_translatable(full_text), label


def _fallback_extract_non_translatable(text: str) -> list:
    """Regex-based fallback extraction when the LLM analysis response cannot be parsed."""
    items: set = set()
    for m in re.finditer(r"1[\s-]?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}", text):
        items.add(m.group())
    for m in re.finditer(r"[\w.+-]+@[\w.-]+\.\w+", text):
        items.add(m.group())
    for m in re.finditer(r"https?://[^\s,)]+", text):
        items.add(m.group())
    for org in ["Molina Healthcare", "BadgerCare Plus", "Molina"]:
        if org in text:
            items.add(org)
    return list(items)


# ---------------------------------------------------------------------------
# Translation Helpers
# ---------------------------------------------------------------------------

def _is_translatable_text(text: str, do_not_translate: list) -> bool:
    """Return True if this text is worth sending to the translation model."""
    stripped = text.strip()
    if not stripped or len(stripped) <= 1:
        return False
    if all(c in "0123456789.,;:!?()-\u2013\u2014/\\|#@$%^&*+=<>{}[]\"'`~ \t\n\r" for c in stripped):
        return False
    if _is_only_non_translatable(stripped, do_not_translate):
        return False
    return True


def _is_only_non_translatable(text: str, do_not_translate: list) -> bool:
    """Return True if text is entirely composed of protected items + punctuation."""
    remaining = text.strip()
    if not remaining:
        return True
    for item in do_not_translate:
        remaining = remaining.replace(item, "")
    remaining = remaining.strip()
    if not remaining:
        return True
    return all(
        c in ".,;:!?()-\u2013\u2014/\\|#@$%^&*+=<>{}[]\"'`~ \t\n\r0123456789"
        for c in remaining
    )


def _apply_translation_to_paragraph(para, translated_text: str) -> None:
    """
    Write translated text into the paragraph's run structure.
    Puts all text into the first run; clears the rest (preserving formatting elements).
    """
    if not para.runs:
        return
    para.runs[0].text = translated_text
    for run in para.runs[1:]:
        run.text = ""


# ---------------------------------------------------------------------------
# Batch Translation
# ---------------------------------------------------------------------------

def _call_batch(
    batch: list,
    target_language: str,
    do_not_translate: list,
    router: TranslationRouter,
    stats: _TranslationStats,
) -> list:
    """
    Translate a batch of short paragraphs in a single API call.
    batch items: (index, paragraph, text). Returns translated strings, same order.
    Single-item batches are translated directly (no numbered markers needed).
    """
    originals = [text for _, _, text in batch]

    if len(batch) == 1:
        _, _, text = batch[0]
        sys_prompt = _build_translation_prompt(target_language, do_not_translate, text)
        result, label = router.call_with_fallback("translate", sys_prompt, text)
        stats.record(label)
        return [result if result else text]

    user_prompt = "\n".join(
        f"[{idx}] {text}" for idx, (_, _, text) in enumerate(batch, 1)
    )
    sys_prompt = _build_batch_system_prompt(target_language, do_not_translate)
    result, label = router.call_with_fallback("translate", sys_prompt, user_prompt)
    stats.batched_calls += 1
    stats.record(label)

    if not result:
        return originals
    return _parse_batch_response(result, len(batch), originals)


def _cache_key(
    text: str,
    target_language: str,
    spanish_target: Optional[SpanishTargetBand],
) -> tuple:
    """Build a cache key that separates Spanish adaptations by target band."""
    base = (text.strip(), target_language)
    if target_language.strip().lower() != "spanish" or not spanish_target:
        return base
    return base + (f"spanish-target:{spanish_target.label}",)


def _maybe_adapt_spanish_text(
    translated_text: str,
    target_language: str,
    do_not_translate: list,
    router: TranslationRouter,
    spanish_target: Optional[SpanishTargetBand],
    spanish_tracker: Optional[SpanishReadabilityTracker],
) -> str:
    """Adapt translated text only for Spanish documents when enabled."""
    if not config.SPANISH_READABILITY_ENABLED:
        return translated_text
    if target_language.strip().lower() != "spanish" or not spanish_target:
        return translated_text
    if not translated_text or not translated_text.strip():
        return translated_text

    result = adapt_spanish_text_to_grade(
        text=translated_text,
        target=spanish_target,
        router=router,
        do_not_translate=do_not_translate,
        max_attempts=config.SPANISH_MAX_ADAPT_PASSES,
        min_words_for_adapt=config.SPANISH_MIN_WORDS_FOR_ADAPT,
    )
    if spanish_tracker:
        spanish_tracker.record(result)
    return result.text


def _print_spanish_readability_summary(summary: Optional[dict]) -> None:
    """Print aggregated Spanish readability/adaptation metrics."""
    if not summary:
        return

    ind = "  "
    print(f"{ind}Spanish readability target : {summary['target_label']}")
    print(
        f"{ind}Target score range        : {summary['target_range']} "
        f"({summary['target_estimated_grade']})"
    )
    print(f"{ind}Fragments processed       : {summary['fragments']}")
    print(f"{ind}Fragments adapted        : {summary['adapted_fragments']}")
    print(f"{ind}Adaptation attempts      : {summary['attempts_total']}")
    print(f"{ind}Avg score before adapt   : {summary['avg_before_score']}")
    print(f"{ind}Avg score after adapt    : {summary['avg_after_score']}")
    print(
        f"{ind}Avg interpreted level    : {summary['avg_after_band']} "
        f"({summary['avg_after_estimated_grade']})"
    )
    print(f"{ind}Target hit on average    : {summary['target_reached_on_average']}")


# ---------------------------------------------------------------------------
# Section-Level Translation
# ---------------------------------------------------------------------------

def _translate_section(
    paragraphs: list,
    target_language: str,
    do_not_translate: list,
    router: TranslationRouter,
    cache: dict,
    stats: _TranslationStats,
    spanish_target: Optional[SpanishTargetBand] = None,
    spanish_tracker: Optional[SpanishReadabilityTracker] = None,
) -> None:
    """
    Translate a list of paragraphs (body / table cell / header / footer).

    Flow:
      1. Collect paragraphs that actually need translation.
      2. Serve cached hits instantly (zero API tokens).
      3. Group short uncached paragraphs -> single batch API call.
      4. Translate long uncached paragraphs individually.
      5. Write all translations back into the paragraph run structure.
    """
    translatable: list = []
    for i, para in enumerate(paragraphs):
        if _paragraph_has_shapes(para):
            continue
        text = extract_paragraph_text(para)
        if _is_translatable_text(text, do_not_translate):
            translatable.append((i, para, text))

    if not translatable:
        return

    result_map: dict = {}
    pending_short: list = []
    pending_long: list = []

    for i, para, text in translatable:
        key = _cache_key(text, target_language, spanish_target)
        if key in cache:
            result_map[i] = cache[key]
            stats.cache_hits += 1
        elif len(text.strip()) < config.BATCH_THRESHOLD:
            pending_short.append((i, para, text))
        else:
            pending_long.append((i, para, text))

    for start in range(0, len(pending_short), config.BATCH_SIZE):
        batch = pending_short[start : start + config.BATCH_SIZE]
        results = _call_batch(batch, target_language, do_not_translate, router, stats)
        for (i, para, text), translated in zip(batch, results):
            translated = _maybe_adapt_spanish_text(
                translated,
                target_language,
                do_not_translate,
                router,
                spanish_target,
                spanish_tracker,
            )
            result_map[i] = translated
            cache[_cache_key(text, target_language, spanish_target)] = translated

    for i, para, text in pending_long:
        sys_prompt = _build_translation_prompt(target_language, do_not_translate, text)
        result, label = router.call_with_fallback("translate", sys_prompt, text)
        translated = result if result else text
        translated = _maybe_adapt_spanish_text(
            translated,
            target_language,
            do_not_translate,
            router,
            spanish_target,
            spanish_tracker,
        )
        result_map[i] = translated
        cache[_cache_key(text, target_language, spanish_target)] = translated
        stats.record(label)

    for i, para, _text in translatable:
        if i in result_map:
            _apply_translation_to_paragraph(para, result_map[i])


# ---------------------------------------------------------------------------
# Document Pipeline
# ---------------------------------------------------------------------------

def _extract_full_document_text(doc) -> str:
    """Extract all visible text from a document for the analysis pass."""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _print_document_stats(
    stats: _TranslationStats,
    source_path: str,
    output_path: str,
    analysis_provider: str,
) -> None:
    total = stats.total_api_calls
    cache_total = total + stats.cache_hits
    savings_pct = round(stats.cache_hits / cache_total * 100) if cache_total > 0 else 0
    ind = "  "
    print(f"\n{ind}DONE!")
    print(f"{ind}Analysis provider  : {analysis_provider}")
    print(f"{ind}Ollama API calls   : {stats.ollama_calls}")
    print(
        f"{ind}Cerebras API calls : {stats.cerebras_calls}"
        + (f" ({stats.fallback_calls} fallback)" if stats.fallback_calls else "")
    )
    print(f"{ind}Batched calls      : {stats.batched_calls}")
    print(f"{ind}Cache hits         : {stats.cache_hits} ({savings_pct}% of requests saved)")
    print(f"{ind}Saved to           : {output_path}")


def translate_document(
    source_path: str,
    output_path: str,
    target_language: str,
    router: TranslationRouter,
    spanish_grade_target: Optional[str] = None,
) -> None:
    """
    Main translation pipeline:
      1. Open .docx and extract all text for analysis.
      2. Agent Step 1 - identify non-translatable elements (Ollama in hybrid mode).
      3. Agent Step 2 - batch-translate body paragraphs (routing by complexity).
      4. Translate table cells.
      5. Translate headers and footers.
      6. Save the translated document.
    """
    print(f"\n{'=' * 70}")
    print(f"Translating  : {os.path.basename(source_path)}")
    print(f"Language     : {target_language}")
    print(f"Routing mode : {config.ROUTING_MODE}")
    print(f"{'=' * 70}")

    doc = Document(source_path)
    full_text = _extract_full_document_text(doc)

    spanish_target: Optional[SpanishTargetBand] = None
    spanish_tracker: Optional[SpanishReadabilityTracker] = None
    if target_language.strip().lower() == "spanish" and config.SPANISH_READABILITY_ENABLED:
        requested_target = spanish_grade_target or config.SPANISH_DEFAULT_TARGET_GRADE
        spanish_target = normalize_target_grade(requested_target)
        spanish_tracker = SpanishReadabilityTracker(spanish_target)
        print(f"  Spanish target readability: {spanish_target.label} ({requested_target})")

    do_not_translate, analysis_provider = analyze_document_content(router, full_text)

    stats = _TranslationStats()
    cache: dict = {}  # per-document phrase cache: {(text_stripped, lang): translation}

    print(f"\n  [Agent Step 2] Translating {len(doc.paragraphs)} paragraphs...")
    _translate_section(
        list(doc.paragraphs), target_language, do_not_translate, router, cache, stats,
        spanish_target=spanish_target, spanish_tracker=spanish_tracker,
    )

    print(f"  Translating {len(doc.tables)} table(s)...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _translate_section(
                    list(cell.paragraphs), target_language,
                    do_not_translate, router, cache, stats,
                    spanish_target=spanish_target, spanish_tracker=spanish_tracker,
                )

    print("  Translating headers and footers...")
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if not hf.is_linked_to_previous:
                _translate_section(
                    list(hf.paragraphs), target_language,
                    do_not_translate, router, cache, stats,
                    spanish_target=spanish_target, spanish_tracker=spanish_tracker,
                )

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)

    _print_document_stats(stats, source_path, output_path, analysis_provider)
    if spanish_tracker:
        _print_spanish_readability_summary(spanish_tracker.summary())
    print(f"{'=' * 70}\n")


# ---------------------------------------------------------------------------
# Legacy API Surface  (backward compatibility)
# ---------------------------------------------------------------------------

def translate_text_block(
    client_or_router,
    text: str,
    target_language: str,
    do_not_translate: list,
) -> str:
    """Translate a single text block. Accepts both old-style client and new router."""
    if not text or not text.strip():
        return text
    stripped = text.strip()
    if all(c in "0123456789.,;:!?()-\u2013\u2014/\\|#@$%^&*+=<>{}[]\"'`~ \t\n\r" for c in stripped):
        return text
    if len(stripped) <= 1:
        return text
    sys_prompt = _build_translation_prompt(target_language, do_not_translate, text)
    from llm_providers import LLMProvider
    if isinstance(client_or_router, TranslationRouter):
        result, _ = client_or_router.call_with_fallback("translate", sys_prompt, text)
    elif isinstance(client_or_router, LLMProvider):
        result = client_or_router.call(sys_prompt, text)
    else:
        result = call_llm(client_or_router, sys_prompt, text)
    if not result:
        print(f"    Warning: Empty translation for '{text[:60]}...' - keeping original.")
        return text
    return result


def _apply_translated_text_to_runs(paragraph, groups, translated_groups):
    """Apply translated text to runs while preserving formatting (legacy helper)."""
    runs = paragraph.runs
    for group, translated_text in zip(groups, translated_groups):
        run_indices = group["run_indices"]
        if not run_indices:
            continue
        first_run_idx = run_indices[0]
        if first_run_idx < len(runs):
            runs[first_run_idx].text = translated_text
        for idx in run_indices[1:]:
            if idx < len(runs):
                runs[idx].text = ""


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="AI Letter Translator - Translate .docx healthcare letters"
    )
    parser.add_argument(
        "--file", type=str, default=None,
        help="Specific filename to translate (from the source folder)",
    )
    parser.add_argument(
        "--all", action="store_true", dest="translate_all",
        help="Translate ALL .docx files in the source folder",
    )
    parser.add_argument(
        "--skip-existing", action="store_true",
        help="Skip files that already exist in the output folder",
    )
    parser.add_argument(
        "--provider",
        choices=["hybrid", "ollama-only", "cerebras-only"],
        default=None,
        help=(
            "LLM routing strategy. "
            "'hybrid' (default): Ollama for simple tasks, Cerebras for complex. "
            "'ollama-only': local model for everything. "
            "'cerebras-only': Cerebras for everything (original behaviour)."
        ),
    )
    parser.add_argument(
        "--spanish-grade",
        type=str,
        default=None,
        help=(
            "Requested Spanish readability target. Accepts numeric grade "
            "(e.g., 6, 8, 10) or band (very easy, easy, normal, moderate, difficult)."
        ),
    )
    args = parser.parse_args()

    if not os.path.exists(config.SOURCE_FOLDER):
        print(f"ERROR: Source folder not accessible: {config.SOURCE_FOLDER}")
        sys.exit(1)

    os.makedirs(config.OUTPUT_FOLDER, exist_ok=True)

    router = TranslationRouter(mode=args.provider)
    print(f"LLM routing  : {args.provider or config.ROUTING_MODE}")

    if args.file:
        files_to_translate = [args.file]
    elif args.translate_all:
        files_to_translate = [
            f for f in os.listdir(config.SOURCE_FOLDER)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]
    else:
        files_to_translate = []
        for f in os.listdir(config.SOURCE_FOLDER):
            if f.lower().endswith(".docx") and not f.startswith("~$"):
                if extract_language_code(f):
                    files_to_translate.append(f)
                    break
        if not files_to_translate:
            print("No files with language codes found in the source folder.")
            sys.exit(1)

    print(f"\nFiles to translate: {len(files_to_translate)}")

    success_count = error_count = 0

    for filename in files_to_translate:
        source_path = os.path.join(config.SOURCE_FOLDER, filename)
        output_path = os.path.join(config.OUTPUT_FOLDER, filename)

        if args.skip_existing and os.path.exists(output_path):
            print(f"  Skipping (already exists): {filename}")
            continue

        lang_code = extract_language_code(filename)
        if not lang_code:
            print(f"  Skipping (no language code found): {filename}")
            continue

        target_language = get_language_name(lang_code)

        try:
            translate_document(
                source_path,
                output_path,
                target_language,
                router,
                spanish_grade_target=args.spanish_grade,
            )
            success_count += 1
        except Exception as exc:
            print(f"  ERROR translating {filename}: {exc}")
            error_count += 1

    print(f"\n{'=' * 70}")
    print("TRANSLATION COMPLETE")
    print(f"  Successful : {success_count}")
    print(f"  Errors     : {error_count}")
    print(f"  Output     : {config.OUTPUT_FOLDER}")
    print(f"{'=' * 70}")




if __name__ == "__main__":
    main()
